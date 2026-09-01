#!/usr/bin/env python3
"""Local PDF/Word ingestion tests."""

import os
import tempfile
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from pypdf import PdfWriter

import sys

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "scripts"))

import distill as cli  # noqa: E402
from file_ingest import article_from_file  # noqa: E402
import dependency_bootstrap as deps  # noqa: E402


def test_pdf_without_text_layer_explains_ocr_requirement():
    with tempfile.TemporaryDirectory() as temp_dir:
        path = Path(temp_dir) / "scan.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=100, height=100)
        with path.open("wb") as handle:
            writer.write(handle)
        try:
            article_from_file(str(path))
        except ValueError as exc:
            assert "没有可提取的文字" in str(exc)
            assert "OCR" in str(exc)
        else:
            raise AssertionError("textless PDF should require OCR")


def test_docx_ingest_preserves_paragraphs_tables_and_overrides():
    with tempfile.TemporaryDirectory() as temp_dir:
        path = os.path.join(temp_dir, "ai-report.docx")
        document = Document()
        document.core_properties.title = "文档元标题"
        document.core_properties.author = "文档作者"
        document.add_paragraph("第一段正文")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "指标"
        table.cell(0, 1).text = "结果"
        table.cell(1, 0).text = "准确率"
        table.cell(1, 1).text = "92%"
        document.add_paragraph("第二段正文")
        document.save(path)

        article = article_from_file(path)
        assert article.title == "文档元标题"
        assert article.author == "文档作者"
        assert article.text.index("第一段正文") < article.text.index("指标") < article.text.index("第二段正文")
        assert "准确率 | 92%" in article.text

        overridden = article_from_file(path, title="覆盖标题", author="覆盖作者")
        assert overridden.title == "覆盖标题"
        assert overridden.author == "覆盖作者"


def test_direct_local_path_and_unsupported_file_message():
    with tempfile.TemporaryDirectory() as temp_dir:
        source = Path(temp_dir) / "article.md"
        source.write_text("# 标题\n\n这是本地文件正文。" * 20, encoding="utf-8")
        args = SimpleNamespace(
            from_text="",
            url=str(source),
            title="",
            author="",
            chart_ocr=False,
        )
        article = cli._get_article(args)
        assert article.text_chars > 100

        path = os.path.join(temp_dir, "notes.rtf")
        with open(path, "w", encoding="utf-8") as handle:
            handle.write("{\\rtf1 不支持}")
        try:
            article_from_file(path)
        except ValueError as exc:
            assert "暂不支持" in str(exc)
            assert ".pdf" in str(exc) and ".docx" in str(exc)
        else:
            raise AssertionError("unsupported file extension was accepted")


def test_dependency_failure_explains_recovery_command():
    original_find_spec = deps.importlib.util.find_spec
    original_run = deps.subprocess.run

    class Failed:
        returncode = 1
        stdout = "permission denied"

    deps.importlib.util.find_spec = lambda _name: None
    deps.subprocess.run = lambda *args, **kwargs: Failed()
    try:
        try:
            deps.ensure_python_dependencies(["pypdf"])
        except RuntimeError as exc:
            message = str(exc)
            assert "自动安装依赖失败" in message
            assert "-m pip install" in message
            assert "pypdf" in message
        else:
            raise AssertionError("missing dependency should fail with an actionable message")
    finally:
        deps.importlib.util.find_spec = original_find_spec
        deps.subprocess.run = original_run


if __name__ == "__main__":
    test_pdf_without_text_layer_explains_ocr_requirement()
    test_docx_ingest_preserves_paragraphs_tables_and_overrides()
    test_direct_local_path_and_unsupported_file_message()
    test_dependency_failure_explains_recovery_command()
    print("file ingest tests passed")
